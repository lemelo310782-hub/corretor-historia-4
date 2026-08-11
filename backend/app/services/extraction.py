"""
Serviço de extração de texto de documentos.

Suporta:
- PDF com camada de texto nativa (extração direta, rápida e sem erros de OCR)
- PDF escaneado / imagem de baixa qualidade (fallback automático para OCR)
- DOCX (parágrafos + tabelas)
- Imagens soltas (PNG/JPG) via OCR

Cada função devolve texto ou lança ValueError com mensagem clara — nunca
falha silenciosamente, porque um professor precisa saber SE e POR QUE
uma ficha não pôde ser lida.
"""
from pathlib import Path

from pypdf import PdfReader
from pdf2image import convert_from_path
import pytesseract
from docx import Document as DocxDocument
from PIL import Image

from app.config import OCR_LANGUAGE, MIN_CHARS_TEXTO_NATIVO


def extrair_texto_pdf(caminho: Path) -> tuple[str, bool]:
    """
    Retorna (texto, usou_ocr).

    Tenta extração nativa primeiro (PDFs gerados digitalmente, ex: exportado
    de um editor de texto). Se o resultado for curto demais — sinal de que é
    um PDF escaneado sem camada de texto — cai para OCR, página por página.
    """
    texto_nativo = ""
    try:
        leitor = PdfReader(str(caminho))
        paginas = [pagina.extract_text() or "" for pagina in leitor.pages]
        texto_nativo = "\n".join(paginas).strip()
    except Exception:
        texto_nativo = ""

    if len(texto_nativo) >= MIN_CHARS_TEXTO_NATIVO:
        return texto_nativo, False

    # Fallback: PDF provavelmente escaneado -> converte páginas em imagens e roda OCR
    try:
        imagens = convert_from_path(str(caminho))
    except Exception as e:
        raise ValueError(
            f"Não foi possível processar o PDF (sem texto nativo e falha ao converter em "
            f"imagem para OCR): {e}. Verifique se o Poppler está instalado — veja o README."
        )

    partes_ocr = []
    for imagem in imagens:
        partes_ocr.append(pytesseract.image_to_string(imagem, lang=OCR_LANGUAGE))
    texto_ocr = "\n".join(partes_ocr).strip()

    if not texto_ocr:
        raise ValueError(
            "Não foi possível extrair texto deste PDF, nem como texto nativo nem via OCR. "
            "Verifique a qualidade do escaneamento (nitidez, contraste, orientação da página)."
        )
    return texto_ocr, True


def extrair_texto_docx(caminho: Path) -> str:
    """Extrai parágrafos e conteúdo de tabelas de um .docx."""
    doc = DocxDocument(str(caminho))
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if celula.text.strip():
                    partes.append(celula.text.strip())

    texto = "\n".join(partes).strip()
    if not texto:
        raise ValueError("O documento DOCX foi lido, mas não contém texto (arquivo vazio?).")
    return texto


def extrair_texto_imagem(caminho: Path) -> str:
    """OCR direto sobre uma imagem solta (foto de ficha preenchida à mão, por exemplo)."""
    imagem = Image.open(caminho)
    texto = pytesseract.image_to_string(imagem, lang=OCR_LANGUAGE).strip()
    if not texto:
        raise ValueError(
            "Não foi possível reconhecer texto nesta imagem. Tente uma foto mais nítida, "
            "bem iluminada e com a página inteira enquadrada."
        )
    return texto


def extrair_texto(caminho_absoluto: Path) -> tuple[str, bool]:
    """
    Dispatcher por extensão de arquivo.
    Retorna (texto_extraido, usou_ocr).
    """
    ext = caminho_absoluto.suffix.lower()
    if ext == ".pdf":
        return extrair_texto_pdf(caminho_absoluto)
    if ext == ".docx":
        return extrair_texto_docx(caminho_absoluto), False
    if ext in (".png", ".jpg", ".jpeg"):
        return extrair_texto_imagem(caminho_absoluto), True

    raise ValueError(f"Extensão não suportada para extração de texto: '{ext}'")
